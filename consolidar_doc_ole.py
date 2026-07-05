import os
import re
import sys
import olefile
from docx import Document

# Configurar encoding para consola Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

DOWNLOADS_DIR = r"output\descargas"
MASTER_DOC_PATH = r"output\digesto_consolidado_word.docx"

if not os.path.exists(DOWNLOADS_DIR):
    print(f"Error: La carpeta {DOWNLOADS_DIR} no existe.")
    sys.exit(1)

# 1. Buscar y ordenar los archivos .doc antiguos
doc_files = [f for f in os.listdir(DOWNLOADS_DIR) if f.lower().endswith(".doc")]
if not doc_files:
    print("No se encontraron archivos .doc para consolidar.")
    sys.exit(0)

def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

doc_files.sort(key=natural_sort_key)
print(f"Se encontraron {len(doc_files)} archivos .doc antiguos para consolidar.")

# 2. Cargar el documento maestro existente de Word
if os.path.exists(MASTER_DOC_PATH):
    print(f"Cargando documento maestro existente: {MASTER_DOC_PATH}")
    doc_master = Document(MASTER_DOC_PATH)
else:
    print("No se encontro un documento maestro de Word previo. Creando uno nuevo...")
    doc_master = Document()
    doc_master.add_heading("Digesto Municipal Consolidado (DOC)", 0)

# Agregar separador de seccion para los documentos antiguos
doc_master.add_page_break()
heading_sec = doc_master.add_heading(level=1)
run_sec = heading_sec.add_run("SECCIÓN: DOCUMENTOS HISTÓRICOS (.DOC)")
run_sec.bold = True
doc_master.add_page_break()

# Limpiar texto para que sea 100% compatible con XML (lxml/python-docx)
def limpiar_para_xml(text):
    if not text:
        return ""
    # XML 1.0 permite: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    # Filtramos cualquier caracter fuera de este rango
    illegal_chars_re = re.compile(
        u'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff\ufeff\ufffe\uffff]'
    )
    clean_text = illegal_chars_re.sub('', text)
    # Reemplazar cualquier nulo restante por seguridad
    clean_text = clean_text.replace('\x00', '')
    return clean_text

# 3. Funcion para extraer texto plano de un OLE .doc de Word
def extraer_texto_doc_ole(filepath):
    try:
        if not olefile.isOleFile(filepath):
            return ""
        
        with olefile.OleFileIO(filepath) as ole:
            if ole.exists('WordDocument'):
                data = ole.openstream('WordDocument').read()
                
                text = ""
                try:
                    text = data.decode('utf-16-le', errors='ignore')
                except Exception:
                    pass
                
                if not text or len(text.strip()) < 50:
                    text = data.decode('latin-1', errors='ignore')
                
                # Normalizar saltos de linea
                clean_text = text.replace('\r\n', '\n').replace('\r', '\n')
                
                # Truncar lineas y quedarnos con lo legible
                lines = clean_text.split('\n')
                filtered_lines = []
                for line in lines:
                    line_strip = line.strip()
                    # Ignorar lineas de basura binaria cortas
                    if len(line_strip) > 2:
                        # Si tiene palabras o numeros
                        if any(c.isalnum() for c in line_strip):
                            # Limpiar caracteres ilegales XML de esta linea
                            xml_line = limpiar_para_xml(line_strip)
                            if xml_line.strip():
                                filtered_lines.append(xml_line.strip())
                
                return "\n".join(filtered_lines)
        return ""
    except Exception as e:
        print(f"    Error extrayendo texto: {e}")
        return ""

# 4. Iterar y agregar al documento maestro
consolidados = 0

for idx, filename in enumerate(doc_files):
    filepath = os.path.join(DOWNLOADS_DIR, filename)
    print(f"[{idx+1}/{len(doc_files)}] Extrayendo y consolidando: {filename}...")
    
    texto = extraer_texto_doc_ole(filepath)
    if texto.strip():
        # Agregar titulo destacado de la norma
        heading = doc_master.add_heading(level=2)
        run = heading.add_run(filename.replace(".doc", ""))
        run.font.name = "Arial"
        
        # Agregar parrafos de la norma
        paragraphs = texto.split("\n")
        for para in paragraphs:
            para_strip = para.strip()
            if para_strip and len(para_strip) > 3:
                try:
                    doc_master.add_paragraph(para_strip)
                except Exception as e:
                    # Si falla por algun caracter raro persistente, lo limpiamos de forma extrema
                    extremo = "".join(c for c in para_strip if ord(c) >= 32 and ord(c) != 127)
                    if extremo.strip():
                        doc_master.add_paragraph(extremo.strip())
                
        # Salto de pagina entre normas
        doc_master.add_page_break()
        consolidados += 1
    else:
        print(f"    Advertencia: No se pudo extraer texto legible de {filename}")

# 5. Guardar el archivo master unificado completo
if consolidados > 0:
    print(f"\nGuardando documento unificado completo en {MASTER_DOC_PATH}...")
    try:
        doc_master.save(MASTER_DOC_PATH)
        print(f"¡Éxito! Se consolidaron {consolidados} archivos .doc adicionales en el maestro.")
    except Exception as e:
        print(f"Error al guardar el archivo de salida: {e}")
else:
    print("No se pudo consolidar ningun archivo .doc.")
