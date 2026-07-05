import os
import re
import sys
from docx import Document

# Configurar encoding para consola Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

DOWNLOADS_DIR = r"output\descargas"
OUTPUT_FILE = r"output\digesto_consolidado_word.docx"

if not os.path.exists(DOWNLOADS_DIR):
    print(f"Error: La carpeta {DOWNLOADS_DIR} no existe.")
    sys.exit(1)

# 1. Obtener y ordenar los archivos .docx de forma natural
docx_files = [f for f in os.listdir(DOWNLOADS_DIR) if f.lower().endswith(".docx")]
if not docx_files:
    print("No se encontraron archivos .docx en la carpeta de descargas.")
    sys.exit(0)

# Ordenamiento natural (ej: 2 antes que 10)
def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

docx_files.sort(key=natural_sort_key)
print(f"Se encontraron {len(docx_files)} archivos .docx para consolidar.")

# 2. Crear el documento base consolidado
doc_master = Document()
doc_master.add_heading("Digesto Municipal Consolidado", 0)
doc_master.add_paragraph("Compilación de ordenanzas, decretos y resoluciones en formato editable.")
doc_master.add_page_break()

# 3. Función para copiar contenido de un documento a otro
def append_document(master, sub_doc_path, doc_name):
    try:
        sub_doc = Document(sub_doc_path)
        
        # Agregar titulo destacado de la norma
        heading = master.add_heading(level=1)
        run = heading.add_run(doc_name.replace(".docx", ""))
        run.font.name = "Arial"
        
        # Copiar todos los parrafos
        for para in sub_doc.paragraphs:
            # Solo copiar si tiene texto para evitar exceso de espacios en blanco
            if para.text.strip():
                new_para = master.add_paragraph()
                # Copiar los "runs" para preservar formato basico (negrita, cursiva)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    
        # Salto de pagina para separar del siguiente documento
        master.add_page_break()
        return True
    except Exception as e:
        print(f"  -> Error al copiar {doc_name}: {e}")
        return False

# 4. Iterar y consolidar
consolidados = 0
for idx, filename in enumerate(docx_files):
    filepath = os.path.join(DOWNLOADS_DIR, filename)
    print(f"[{idx+1}/{len(docx_files)}] Consolidando: {filename}...")
    if append_document(doc_master, filepath, filename):
        consolidados += 1

# 5. Guardar el archivo master consolidado
print(f"\nGuardando documento unificado en {OUTPUT_FILE}...")
try:
    doc_master.save(OUTPUT_FILE)
    print(f"¡Éxito! Se consolidaron {consolidados} archivos en: {OUTPUT_FILE}")
except Exception as e:
    print(f"Error al guardar el archivo de salida: {e}")
