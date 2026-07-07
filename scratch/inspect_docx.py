import docx

doc = docx.Document("output/digesto_consolidado_word.docx")
print("Total párrafos:", len(doc.paragraphs))

# Mostrar encabezados y párrafos clave de inicio de normas
heading_count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if p.style.name.startswith("Heading"):
        print(f"[Heading {i}] {p.style.name} -> {text}")
        heading_count += 1
        if heading_count >= 15:
            break

print("\n--- Párrafos con ORDENANZA N o similar ---")
norma_count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    upper = text.upper()
    if ("ORDENANZA" in upper or "DECRETO" in upper or "RESOLUCI" in upper) and ("N" in upper or "N°" in upper or "Nº" in upper or "NUMERO" in upper):
        print(f"[Norma text {i}] -> {text[:150]}")
        norma_count += 1
        if norma_count >= 15:
            break

